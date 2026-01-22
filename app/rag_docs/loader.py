# 5. app/ingestion/loader.py
# 下面的代码从一个目录下递归加载各种文档文件pdf/docx/doc/md/txt等，将它们读成LangChain的 Document对象列表。然后用RecursiveCharacterTextSplitter把这些文档按在settings里配置的
# chunk_size和chunk_overlap切成小块，方便后面做向量化/检索。
#
# chunk表示切出来的一小块文本。
#
# chunk_size是每个文本块的最大长度（按字符数算）。
# 块太大:
# 	向量长度大，embedding变慢、占空间
# 	模型一次看太多文本，反而不精细
# 块太小:
# 	语义被切碎，一块的信息太少
# 	需要更多块才能覆盖同一篇文档
#
# 典型取值512~1500字符，如果要语义更完整1000左右比较常见
# 如果文档都特别长想加速处理可以用1500–2000
#
# chunk_overlap是相邻两个chunk之间的重叠字符数。
# 比如chunk_size = 1000，chunhostname -I
# k_overlap = 200
#
# 切出来是这样的
# 第1块：0 ~ 999
# 第2块：800 ~ 1799
# 第3块：1600 ~ 2599
#
# 为什么要重叠？
# 文本语义经常跨chunk，比如：段落结尾一句话 + 下一段开头一起才好理解
# 如果没有重叠，检索某些句子时，可能刚好被切开，导致模型missing context
# 有重叠可以让模型在任意块里都看到一部分上下文
#
# 一般是chunk_size的 10%～30%
# 如果取值太大如800会：增加很多重复内容、向量库变大，检索变慢

from pathlib import Path
from typing import List
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
import docx
from app.config import settings

def load_pdf(path: Path) -> List[Document]:
	# """下面的代码将pdf分割成单独的页面，每一个页面的文本被封装成一个Document放入list"""
	reader = PdfReader(str(path))
	docs = []
	for i, page in enumerate(reader.pages):
		text = page.extract_text() or ""
		if text.strip():
			docs.append(Document(
				page_content=text,
				metadata={"source": str(path), "page": i+1}
			))
	return docs

def load_docx(path: Path) -> List[Document]:
	d = docx.Document(str(path))
	text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
	return [Document(page_content=text, metadata={"source": str(path)})] if text else []

# def load_text(path: Path) -> List[Document]:
#
# 	return docs.append(Document(page_content=f.read_text(encoding="utf-8"),
# 								 metadata={"source": str(f)}))


def load_docs(dir_path: str) -> List[Document]:
	p = Path(dir_path)
	docs: List[Document] = []
	for f in p.rglob("*"):
		if f.suffix.lower() == ".pdf":
			docs.extend(load_pdf(f))
		elif f.suffix.lower() in [".docx", ".doc"]:
			docs.extend(load_docx(f))
		elif f.suffix.lower() in [".md", ".txt"]:
			docs.append(Document(page_content=f.read_text(encoding="utf-8"),
								 metadata={"source": str(f)}))
	return docs

def split_docs(docs: List[Document]) -> List[Document]:
	splitter = RecursiveCharacterTextSplitter(
		chunk_size=settings.chunk_size,
		chunk_overlap=settings.chunk_overlap
	)
	return splitter.split_documents(docs)

# 第一个函数主要是为了应对后续文件单独追加而设计的，他就是处理一个单独的文件而已。
def load_single_file(path: Path) -> List[Document]:
    """根据文件后缀加载文件，返回LangChain的 Document列表"""
    suf = path.suffix.lower()
    if suf == ".pdf":
        return load_pdf(path)
    if suf in [".docx", ".doc"]:
        return load_docx(path)
    if suf in [".md", ".txt"]:
        text = path.read_text(encoding="utf-8")
        return [Document(page_content=text, metadata={"source": str(path)})] if text.strip() else []
    return []

# 第二个函数主要是把一批Document切成小块，并且给每一小块贴上权限标签visibility和文档ID。
def split_with_visibility(
    docs: List[Document],
    visibility: str,
    doc_id: str | None = None,
    extra_meta: dict | None = None,
) -> List[Document]:
    chunks = split_docs(docs)
    extra_meta = dict(extra_meta or {})  # 扩展元数据（参数中额外提供）
    for c in chunks:  # 循环的是切割好的每一个小块
        c.metadata = dict(c.metadata or {}) # 每一个小块固有的元数据取出来
        c.metadata["visibility"] = visibility  # 加入可见性元数据
        if doc_id: # 如果手动提供了文档id也加为元数据
            c.metadata["doc_id"] = doc_id
        for k, v in extra_meta.items():
            if v is not None:
                c.metadata[k] = v  # 将参数提供的额外元数据也加进来
    return chunks # 事无巨细的没有遗漏任何元数据，还给了动态添加元数据的机会extra_meta



if __name__ == "__main__":
	for l in load_docs('../data'):
		print('-------------begin--------------')
		print(l)
		print('--------------end-------------')
